from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

document = Document()

# Configuración básica de estilo
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)

# Título principal
document.add_heading('Derivadas Comunes', 0)

# Sección 1: Derivadas Básicas
document.add_heading('1. Derivadas Básicas', level=1)
table = document.add_table(rows=5, cols=2)
table.style = 'Table Grid'
table.autofit = True

# Encabezados de tabla
row = table.rows[0]
row.cells[0].text = 'Función'
row.cells[1].text = 'Derivada'

# Datos de la tabla
data_basicas = [
    ('f(x) = c', 'f\'(x) = 0'),
    ('f(x) = x^n', 'f\'(x) = n·x^{n-1}'),
    ('f(x) = e^x', 'f\'(x) = e^x'),
    ('f(x) = ln(x)', 'f\'(x) = 1/x'),
]

for i, (func, deriv) in enumerate(data_basicas, start=1):
    row = table.rows[i]
    row.cells[0].text = func
    row.cells[1].text = deriv

# Sección 2: Derivadas Trigonométricas
document.add_heading('2. Derivadas de Funciones Trigonométricas', level=1)
table = document.add_table(rows=7, cols=2)
table.style = 'Table Grid'

row = table.rows[0]
row.cells[0].text = 'Función'
row.cells[1].text = 'Derivada'

data_trig = [
    ('sin(x)', 'cos(x)'),
    ('cos(x)', '-sin(x)'),
    ('tan(x)', 'sec²(x)'),
    ('cot(x)', '-csc²(x)'),
    ('sec(x)', 'sec(x)tan(x)'),
    ('csc(x)', '-csc(x)cot(x)'),
]

for i, (func, deriv) in enumerate(data_trig, start=1):
    row = table.rows[i]
    row.cells[0].text = func
    row.cells[1].text = deriv

# Sección 3: Derivadas Exponenciales/Logarítmicas
document.add_heading('3. Derivadas de Funciones Exponenciales y Logarítmicas', level=1)
table = document.add_table(rows=3, cols=2)
table.style = 'Table Grid'

row = table.rows[0]
row.cells[0].text = 'Función'
row.cells[1].text = 'Derivada'

data_exp_log = [
    ('a^x', 'a^x · ln(a)'),
    ('log_a(x)', '1 / (x·ln(a))'),
]

for i, (func, deriv) in enumerate(data_exp_log, start=1):
    row = table.rows[i]
    row.cells[0].text = func
    row.cells[1].text = deriv

# Sección 4: Funciones Inversas
document.add_heading('4. Derivadas de Funciones Inversas', level=1)
table = document.add_table(rows=4, cols=2)
table.style = 'Table Grid'

row = table.rows[0]
row.cells[0].text = 'Función'
row.cells[1].text = 'Derivada'

data_inversas = [
    ('arcsin(x)', '1 / √(1 - x²)'),
    ('arccos(x)', '-1 / √(1 - x²)'),
    ('arctan(x)', '1 / (1 + x²)'),
]

for i, (func, deriv) in enumerate(data_inversas, start=1):
    row = table.rows[i]
    row.cells[0].text = func
    row.cells[1].text = deriv

# Sección 5: Reglas de Derivación
document.add_heading('5. Reglas de Derivación', level=1)

reglas = [
    "Regla de la Cadena: d/dx [f(g(x))] = f'(g(x)) · g'(x)",
    "Regla del Producto: d/dx [f(x)g(x)] = f'(x)g(x) + f(x)g'(x)",
    "Regla del Cociente: d/dx [f(x)/g(x)] = [f'(x)g(x) - f(x)g'(x)] / [g(x)]²"
]

for regla in reglas:
    p = document.add_paragraph(regla, style='List Bullet')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Sección 6: Ejemplos Prácticos
document.add_heading('6. Ejemplos Prácticos', level=1)

ejemplos = [
    ("Ejemplo 1: Derivar f(x) = 3x⁴ + 2sin(x)", 
     "f'(x) = 12x³ + 2cos(x)"),
    
    ("Ejemplo 2: Derivar f(x) = e²ˣ · ln(x)", 
     "f'(x) = 2e²ˣ·ln(x) + e²ˣ/x (Regla del Producto)"),
]

for titulo, solucion in ejemplos:
    p = document.add_paragraph()
    p.add_run(titulo).bold = True
    p.add_run("\n" + solucion)

document.save('derivadas_completas.docx')